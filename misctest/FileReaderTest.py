import os
import chardet
import pandas as pd
import docx
import pdfplumber
import hashlib
import io


class TextReader:
    def __init__(self, s: str):
        self.__text = s
        self.__text_lines = s.splitlines()

    @property
    def text(self):
        return self.__text

    @property
    def text_lines(self):
        return self.__text_lines

    def get_value(self, key, sep=None):
        index = self.__text.find(key)
        if index >= 0:
            return self.__text[index + len(key):].lstrip().split(sep)[0].strip()

    def index_of(self, *args, bool_func=all):
        for i in range(len(self.text_lines)):
            if bool_func([arg in self.text_lines[i] for arg in args]):
                return i
        return -1


class FileReader:
    def __init__(self, fn):
        self._filename = fn
        self._basename = os.path.basename(self._filename)
        self._dirname = os.path.dirname(self._filename)
        self._clear_name, self._suffix = os.path.splitext(self._basename)
        with open(self.filename,  'rb') as f:
            self._buffer = f.read()
        self._md5 = hashlib.md5(self._buffer).hexdigest()

    @property
    def filename(self):
        return self._filename

    @property
    def basename(self):
        return self._basename

    @property
    def dirname(self):
        return self._dirname

    @property
    def clear_name(self):
        return self._clear_name

    @property
    def suffix(self):
        return self._suffix

    @property
    def md5(self):
        return self._md5

    @property
    def buffer(self):
        return self._buffer


class ExcelReader(FileReader):
    def __init__(self, fn):
        super().__init__(fn)
        self._excel_file = pd.ExcelFile(io.BytesIO(self.buffer))

        class Item:
            def __init__(self, item_func):
                self.__item_func = item_func

            def __getitem__(self, arg):
                return self.__item_func(arg)

        self.sheet = Item(self.__get_sheet)
        self.__sheets = {}

    def __get_sheet(self, sheet_name):
        if sheet_name not in self.__sheets:
            self.__sheets[sheet_name] = self.open(sheet_name)
        return self.__sheets[sheet_name]

    def open(self, sheet_name, header=None) -> pd.DataFrame:
        ret = self._excel_file.parse(sheet_name, header)

        def num_to_letters(num):
            letters = []
            while num:
                num, remainder = divmod(num - 1, 26)
                letters.append(chr(remainder + ord('A')))
            return ''.join(reversed(letters))

        if header is None:
            ret.columns = [num_to_letters(int(c) + 1) for c in ret.columns]
            ret.index = ret.index + 1

        return ret

    def cut(self, sheet_name, rect=pd.IndexSlice[:, :]):
        ret = self.sheet[sheet_name].loc[rect]
        ret.columns = ret.iloc[0]
        ret = ret.iloc[1:]
        ret = ret.dropna(axis=0, how='all')
        return ret

    @property
    def sheet_names(self):
        return self._excel_file.sheet_names


class TxtReader(FileReader, TextReader):
    def __init__(self, fn):
        super().__init__(fn)
        # with open(fn, 'rb') as f:
        #     data = f.read()
        TextReader.__init__(self, self.buffer.decode(chardet.detect(self.buffer)['encoding']))


class DocxReader(FileReader, TextReader):
    def __init__(self, fn):
        super().__init__(fn)
        self.__document = docx.Document(io.BytesIO(self.buffer))
        TextReader.__init__(self, '\n'.join([p.text for p in self.__document.paragraphs]))
        self.__tables = []
        for t in self.__document.tables:
            table = []
            for r in t.rows:
                table.append([c.text for c in r.cells])
            self.__tables.append(table)

    @property
    def document(self):
        return self.__document

    @property
    def tables(self):
        return self.__tables


class PdfReader(FileReader):
    def __init__(self, fn):
        super().__init__(fn)

        class Page(TextReader):
            def __init__(self, _p):
                # super().__init__(_p.extract_text())
                self.tables = _p.extract_tables()
                self.words_lines = [merge_chars(tl['chars']) for tl in _p.extract_text_lines()]
                super().__init__('\n'.join([' '.join([w['text'] for w in wl]) for wl in self.words_lines]))

        self.__pages = []
        with pdfplumber.open(io.BytesIO(self.buffer)) as f:
            for p in f.pages:
                self.__pages.append(Page(p))

    @property
    def pages(self):
        return self.__pages


def merge_chars(chars):
    """
    合并字符变成单词，其他属性取第一个字符的属性
    :param chars: 字符列表
    :return: 单词列表
    """
    ret = []
    word = chars[0].copy()
    for i in range(1, len(chars)):
        char_a = chars[i - 1]
        char_b = chars[i]
        if char_b['x0'] - char_a['x0'] > char_a['width'] + 1:
            ret.append(word)
            word = chars[i].copy()
        else:
            word['text'] += chars[i]['text']
    ret.append(word)
    return ret

